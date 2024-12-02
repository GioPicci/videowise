from fastapi import FastAPI,HTTPException,BackgroundTasks
from pydantic import BaseModel
import whisperx
import torch
import os
from typing import Optional
from fastapi.responses import StreamingResponse
import httpx
import io
import librosa
import uuid
from threading import Lock
import pdfkit
from htmldocx import HtmlToDocx
from docx import Document
from docx.shared import Pt,RGBColor
import asyncio
import aiohttp
from datetime import timedelta
from bs4 import BeautifulSoup
from pydub import AudioSegment
from threading import Timer

app = FastAPI(title="WhisperX API")

remote_filesystem_url = os.getenv("FILESYSTEM_API_URL") #"http://192.168.202.45:8081"
model_name = os.getenv("WHISPER_MODEL")

print("Filesystem API URL:", remote_filesystem_url)

base_file_path = "/home/badmin/whisperxrestservice/build/classes/java/main/uploads"
audio_dir_path = os.path.join(base_file_path,"extracted_audios")
transcription_dir_path = os.path.join(base_file_path,"transcriptions")

transcription_statuses = {}

model_t_lock = Lock()
model_a_lock = Lock()

class WhisperXModelManager:
    def __init__(self, 
                 model_name="large-v2", 
                 auto_release=True,
                 timeout=300,  # 5 minutes of inactivity before memory is freed
                 device=None):
        
        self.model_name = model_name
        self.auto_release = auto_release
        self.timeout = timeout
        self.device = device or ("cuda" if torch.cuda.is_available() else "cpu")
        
        # Lock per gestire l'accesso concorrente
        self.model_t_lock = Lock()
        self.model_a_lock = Lock()
        
        # Modelli
        self.model_t = None
        self.model_a = None
        self.metadata = None
        
        # Timer per il rilascio
        self.release_timer = None

    def _load_models(self):
        # Caricamento modelli solo se non già caricati
        if not self.model_t:
            print("Loading transcription model")
            self.model_t = whisperx.load_model(
                self.model_name, 
                self.device, 
                compute_type="float16" if self.device == "cuda" else "int8"
            )
        
        if not self.model_a:
            print("Loading alignment model")
            self.model_a, self.metadata = whisperx.load_align_model(
                language_code="it", 
                device=self.device
            )

    def _start_release_timer(self):
        # Annulla il timer precedente se esistente
        if self.release_timer:
            self.release_timer.cancel()
        
        # Imposta un nuovo timer per rilasciare i modelli
        self.release_timer = Timer(self.timeout, self._release_models)
        self.release_timer.start()

    def _release_models(self):
        print("Releasing WhisperX models due to inactivity")
        
        # Rilascio esplicito dei modelli
        if self.model_t:
            del self.model_t
            self.model_t = None
        
        if self.model_a:
            del self.model_a
            self.model_a = None
        
        self.metadata = None
        
        # Pulizia memoria GPU
        torch.cuda.empty_cache()

    def transcribe_and_align(self, audio, language):
        # Load models if necessary
        self._load_models()
        
        try:
            # Transcription with lock
            with self.model_t_lock:
                result = self.model_t.transcribe(audio, language=language)
            
            # Alignment with lock
            with self.model_a_lock:
                result = whisperx.align(
                    result["segments"], 
                    self.model_a, 
                    self.metadata, 
                    audio, 
                    self.device, 
                    return_char_alignments=False
                )
            
            # Restart the release timer if auto_release is active
            if self.auto_release:
                self._start_release_timer()
            
            return result
        
        finally:
            # Always clean the GPU cache
            torch.cuda.empty_cache()


# Load the model once at startup
device = "cuda" if torch.cuda.is_available() else "cpu"

print("WhisperX Models Loading Start")
whisperx_manager = WhisperXModelManager(
    model_name=model_name,
    device="cuda" if torch.cuda.is_available() else "cpu"
)

#model = whisperx.load_model(model_name,device,compute_type="float16" if device == "cuda" else "int8")
#print("Transcription model loading finish, align model start")
#model_a,metadata = whisperx.load_align_model(language_code="it",device=device)
#print("Align model loading finish")


# Define request object
class TranscriptionRequest(BaseModel):
    audio_file: str
    # dst_dir: str
    language: Optional[str] = "it"
    modelName: Optional[str] = "large-v2"


class HTMLContent(BaseModel):
    html: str

async def download_audio(audio_file_url: str) -> io.BytesIO:
    """Asynchronously download audio file."""
    async with httpx.AsyncClient() as client:
        response = await client.get(audio_file_url)
        if response.status_code != 200:
            raise HTTPException(status_code=response.status_code,detail="Could not download file")
        return io.BytesIO(response.content)


def format_timestamp_vtt(seconds: float) -> str:
    """Convert seconds to WebVTT timestamp format HH:MM:SS.mmm"""
    td = timedelta(seconds=seconds)
    hours = td.seconds // 3600
    minutes = (td.seconds % 3600) // 60
    seconds = td.seconds % 60
    milliseconds = round(td.microseconds / 1000)
    return f"{hours:02d}:{minutes:02d}:{seconds:02d}.{milliseconds:03d}"


def create_vtt_content(segments: list) -> str:
    """Convert WhisperX segments to WebVTT format"""
    vtt_lines = ["WEBVTT",""]  # WebVTT header required

    for i,segment in enumerate(segments,1):
        start_time = format_timestamp_vtt(segment['start'])
        end_time = format_timestamp_vtt(segment['end'])
        text = segment['text'].strip()

        vtt_lines.extend([
            f"{i}",
            f"{start_time} --> {end_time}",
            text,
            ""  # Empty line between entries
        ])
    return "\n".join(vtt_lines)


def create_ass_content(segments: list) -> str:
    """Convert WhisperX segments to ASS format with styling"""
    ass_header = """[Script Info]
ScriptType: v4.00+
PlayResX: 384
PlayResY: 288
ScaledBorderAndShadow: yes

[V4+ Styles]
Format: Name, Fontname, Fontsize, PrimaryColour, SecondaryColour, OutlineColour, BackColour, Bold, Italic, Underline, StrikeOut, ScaleX, ScaleY, Spacing, Angle, BorderStyle, Outline, Shadow, Alignment, MarginL, MarginR, MarginV, Encoding
Style: Default,Arial,16,&H00FFFFFF,&H000000FF,&H00000000,&H00000000,0,0,0,0,100,100,0,0,1,2,2,2,10,10,10,1

[Events]
Format: Layer, Start, End, Style, Name, MarginL, MarginR, MarginV, Effect, Text
"""

    def format_time_ass(seconds: float) -> str:
        """Convert seconds to ASS timestamp format H:MM:SS.cc"""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        seconds = seconds % 60
        centiseconds = int((seconds % 1) * 100)
        seconds = int(seconds)
        return f"{hours}:{minutes:02d}:{seconds:02d}.{centiseconds:02d}"

    ass_lines = [ass_header]

    for segment in segments:
        start_time = format_time_ass(segment['start'])
        end_time = format_time_ass(segment['end'])
        text = segment['text'].strip().replace("\n","\\N")

        # Add dialogue line with default style
        dialogue_line = f"Dialogue: 0,{start_time},{end_time},Default,,0,0,0,,{text}"
        ass_lines.append(dialogue_line)

    return "\n".join(ass_lines)


def format_timestamp(seconds: float) -> str:
    """Convert seconds to SRT timestamp format HH:MM:SS,mmm"""
    td = timedelta(seconds=seconds)
    hours = td.seconds // 3600
    minutes = (td.seconds % 3600) // 60
    seconds = td.seconds % 60
    milliseconds = round(td.microseconds / 1000)
    return f"{hours:02d}:{minutes:02d}:{seconds:02d},{milliseconds:03d}"


def create_srt_content(segments: list) -> str:
    """Convert WhisperX segments to SRT format"""
    srt_lines = []
    for i,segment in enumerate(segments,1):
        start_time = format_timestamp(segment['start'])
        end_time = format_timestamp(segment['end'])
        text = segment['text'].strip()

        srt_lines.extend([
            str(i),
            f"{start_time} --> {end_time}",
            text,
            ""  # Empty line between entries
        ])
    return "\n".join(srt_lines)


async def upload_srt_file(srt_content: str,original_filename: str,remote_filesystem_url: str) -> str:
    """Upload SRT file to remote filesystem"""
    srt_filename = original_filename.rsplit('.',1)[0] + '.srt'
    upload_url = f"{remote_filesystem_url}/upload_stream"

    # Create file-like object from string content
    srt_stream = io.BytesIO(srt_content.encode('utf-8'))

    async with aiohttp.ClientSession() as session:
        try:
            async with session.post(
                    upload_url,
                    data=srt_stream,
                    headers={'title': srt_filename}
            ) as response:
                if response.status != 200:
                    raise HTTPException(
                        status_code=500,
                        detail=f"Failed to upload subtitles: {await response.text()}"
                    )
                return srt_filename
        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"Error uploading subtitles: {str(e)}"
            )


async def upload_subtitle_file(content: str,original_filename: str,extension: str,remote_filesystem_url: str) -> str:
    """Upload subtitle file to remote filesystem"""
    subtitle_filename = original_filename.rsplit('.',1)[0] + extension
    upload_url = f"{remote_filesystem_url}/upload_stream"

    # Create file-like object from string content
    subtitle_stream = io.BytesIO(content.encode('utf-8'))

    async with aiohttp.ClientSession() as session:
        try:
            async with session.post(
                    upload_url,
                    data=subtitle_stream,
                    headers={'title': subtitle_filename}
            ) as response:
                if response.status != 200:
                    raise HTTPException(
                        status_code=500,
                        detail=f"Failed to upload subtitles: {await response.text()}"
                    )
                return subtitle_filename
        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"Error uploading subtitles: {str(e)}"
            )


def transcribe_audio_file(audio_bytes: io.BytesIO,format: str,language: str,transcription_id: str) -> dict:
    """Transcribe audio and return result."""
    try:
        print("Format of the file:",format)
        if format == "mp3":
            audio = AudioSegment.from_mp3(audio_bytes)
            audio_bytes = io.BytesIO()
            audio.export(audio_bytes,format="wav")
            audio_bytes.seek(0)
        audio,sr = librosa.load(audio_bytes,sr=16000)
        result = whisperx_manager.transcribe_and_align(audio, language=language)
        #with model_t_lock:
        #    result = model.transcribe(audio,language=language)
        #with model_a_lock:
        #    result = whisperx.align(result["segments"],model_a,metadata,audio,device,return_char_alignments=False)

        # Process result
        processed_result = {
            "segments": [
                {
                    "start": segment["start"],
                    "end": segment["end"],
                    "text": segment["text"]
                } for segment in result["segments"]
            ],
            "language": language
        }

        # Generate subtitle contents
        srt_content = create_srt_content(processed_result["segments"])
        vtt_content = create_vtt_content(processed_result["segments"])
        ass_content = create_ass_content(processed_result["segments"])

        transcription_statuses[transcription_id] = {
            "status": "success",
            "output": result,
            "srt_content": srt_content,
            "vtt_content": vtt_content,
            "ass_content": ass_content
        }
    except Exception as e:
        transcription_statuses[transcription_id] = {"status": "fail"}
        print("Errore trascrizione:",str(e))
        return {
            "status": "fail",
            "message": f"Transcription failed: {str(e)}"
        }

import traceback
import logging

logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

@app.post("/transcribe")
async def transcribe_audio(request: TranscriptionRequest,background_tasks: BackgroundTasks):
    # Start transcription and store a pending status
    transcription_id = str(uuid.uuid4())
    transcription_statuses[transcription_id] = {"status": "processing"}
    audio_file_url = f"{remote_filesystem_url}/download/{request.audio_file}"
    logger.info("Received transcription request: %s", request)
    logger.debug("Generated transcription ID: %s", transcription_id)
    logger.debug("Audio file URL: %s", audio_file_url)

    # Download the audio file asynchronously
    try:
        logger.info("Attempting to download audio file from URL: %s", audio_file_url)
        transcription_statuses[transcription_id]["audio_bytes"] = await download_audio(audio_file_url)
        logger.info("Audio file downloaded successfully.")
    except HTTPException as e:
        logger.error("Error downloading audio file: %s", e)
        logger.debug("Traceback:\n%s", traceback.format_exc())
        transcription_statuses[transcription_id]["status"] = "fail"
        return {"status": "fail", "message": "Could not download audio file"}
    except Exception as e:
        logger.error("Unexpected error during audio file download: %s", e)
        logger.debug("Traceback:\n%s", traceback.format_exc())
        transcription_statuses[transcription_id]["status"] = "fail"
        return {"status": "fail", "message": "Unexpected error during audio file download"}

    # Log request data
    logger.debug("Request data: %s", request)
    logger.debug("Language for transcription: %s", request.language)

    # Schedule transcription as a background task
    try:
        logger.info("Scheduling transcription task.")
        background_tasks.add_task(
            transcribe_audio_file,
            transcription_statuses[transcription_id]["audio_bytes"],
            os.path.basename(request.audio_file),
            request.language,
            transcription_id
        )
    except Exception as e:
        logger.error("Error scheduling transcription task: %s", e)
        logger.debug("Traceback:\n%s", traceback.format_exc())
        transcription_statuses[transcription_id]["status"] = "fail"
        return {"status": "fail", "message": "Error scheduling transcription task"}

    # Schedule SRT upload as a separate background task
    async def upload_subtitles_when_ready():
        while True:
            try:
                status = transcription_statuses.get(transcription_id, {})
                logger.debug("Checking transcription status: %s", status)
                if status.get("status") == "success":
                    try:
                        logger.info("Uploading subtitles for transcription ID: %s", transcription_id)
                        srt_filename = await upload_subtitle_file(
                            status["srt_content"], request.audio_file, ".srt", remote_filesystem_url
                        )
                        status["srt_filename"] = srt_filename

                        vtt_filename = await upload_subtitle_file(
                            status["vtt_content"], request.audio_file, ".vtt", remote_filesystem_url
                        )
                        status["vtt_filename"] = vtt_filename

                        ass_filename = await upload_subtitle_file(
                            status["ass_content"], request.audio_file, ".ass", remote_filesystem_url
                        )
                        status["ass_filename"] = ass_filename

                        # Clean up after upload
                        del status["srt_content"]
                        del status["vtt_content"]
                        del status["ass_content"]
                        logger.info("Subtitle files uploaded successfully.")
                    except Exception as e:
                        logger.error("Error uploading subtitle files: %s", e)
                        logger.debug("Traceback:\n%s", traceback.format_exc())
                    break
                elif status.get("status") == "fail":
                    logger.error("Transcription failed for ID: %s", transcription_id)
                    break
                await asyncio.sleep(1)  # Check every second
            except Exception as e:
                logger.error("Unexpected error during subtitle upload: %s", e)
                logger.debug("Traceback:\n%s", traceback.format_exc())
                break

    try:
        logger.info("Scheduling subtitle upload task.")
        background_tasks.add_task(upload_subtitles_when_ready)
    except Exception as e:
        logger.error("Error scheduling subtitle upload task: %s", e)
        logger.debug("Traceback:\n%s", traceback.format_exc())
        transcription_statuses[transcription_id]["status"] = "fail"
        return {"status": "fail", "message": "Error scheduling subtitle upload task"}

    # Immediate response
    logger.info("Transcription task initiated successfully. Returning response.")
    return {
        "transcription_token": transcription_id,
        "status": "processing",
        "message": "Transcription is being processed in the background."
    }


@app.get("/check_status/{transcription_id}")
async def check_status(transcription_id: str):
    status = transcription_statuses.get(transcription_id)
    if not status:
        raise HTTPException(status_code=404,detail="Transcription ID not found")
    return status


@app.get("/retrieve_output/{transcription_id}")
async def retrieve_output(transcription_id: str):
    status = transcription_statuses.get(transcription_id)
    if not status:
        raise HTTPException(status_code=404,detail="Transcription ID not found")

    output = status["output"]
    if output is None:
        raise HTTPException(status_code=204,detail="Transcription not ready yet")
    transcription_statuses.pop(transcription_id)
    return {
        "status": "success",
        "output": output
    }


@app.post("/html_to_pdf")
async def convert_html_to_pdf(content: HTMLContent):
    try:
        print("Requested endpoint /html_to_pdf")
        html_with_meta = f"""<!DOCTYPE html>
        <html lang="en">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>Document</title>
            <style>
                body {{
                    margin: 0;
                    padding: 0;
                    font-family: Arial, sans-serif;
                }}

                /* Rule for all elements after h2 */
                h2 ~ *:not(h2) {{
                    margin-left: 15px;
                }}

                /* Rule for all elements after h3 */
                h3 ~ *:not(h3, h2) {{
                    margin-left: 30px;
                }}

                /* Rule for all elements after h4 */
                h4 ~ *:not(h4, h3, h2) {{
                    margin-left: 45px;
                }}

                /* Rule for all elements after h5 */
                h5 ~ *:not(h5, h4, h3, h2) {{
                    margin-left: 60px;
                }}

                /* Rule for all elements after h6 */
                h6 ~ *:not(h6, h5, h4, h3, h2) {{
                    margin-left: 75px;
                }}

                h2,h3,h4,h5,h6 {{
                    margin-bottom: 0px;
                }}
            </style>
        </head>
        <body>
            {content.html}
        </body>
        </html>
        """
        print("HTML RICEVUTO:", content.html)
        # Convert HTML to PDF
        pdf = pdfkit.from_string(html_with_meta,False)

        # Create a BytesIO object to hold the PDF
        pdf_io = io.BytesIO(pdf)
        pdf_io.seek(0)

        return StreamingResponse(pdf_io,media_type='application/pdf',
                                 headers={"Content-Disposition": "attachment; filename=output.pdf"})
    except Exception as e:
        raise HTTPException(status_code=500,detail=str(e))


def apply_heading_style(heading_paragraph,level):
    """Apply proper font size and styling to headings"""
    run = heading_paragraph.runs[0] if heading_paragraph.runs else heading_paragraph.add_run()

    # Define heading sizes (in points)
    heading_sizes = {
        1: 24,  # H1
        2: 18,  # H2
        3: 14,  # H3
        4: 12,  # H4
        5: 10,  # H5
        6: 10  # H6
    }

    # Apply size
    run.font.size = Pt(heading_sizes.get(level,12))

    # Make heading bold
    run.bold = True

    # Add some spacing before the heading
    heading_paragraph.paragraph_format.space_before = Pt(14)
    heading_paragraph.paragraph_format.space_after = Pt(4)

    # Optional: Add heading color
    run.font.color.rgb = RGBColor(0,0,0)  # Black color


def html_markdown_word(html_content,filename):
    # Create new Document
    doc = Document()

    # Parse HTML
    soup = BeautifulSoup(html_content,'html.parser')

    def process_element(element):
        """Process individual HTML elements"""
        if element.name and element.name.startswith('h') and len(element.name) == 2:
            # Handle all heading levels (h1-h6)
            level = int(element.name[1])
            heading = doc.add_paragraph()
            heading.add_run(element.get_text().strip())
            apply_heading_style(heading,level)

        elif element.name == 'div':
            # Process the content within div
            p = doc.add_paragraph()
            process_inline_elements(element,p)

        elif element.name == 'br':
            doc.add_paragraph()  # Add empty paragraph for line break

        elif element.name == 'ul':
            for li in element.find_all('li',recursive=False):
                p = doc.add_paragraph(style='List Bullet')
                process_inline_elements(li,p)

        elif element.name == 'strong':
            return {'text': element.get_text(),'bold': True}

        elif element.name == 'em':
            return {'text': element.get_text(),'italic': True}

        else:
            # Handle text nodes
            text = element.string
            if text and text.strip():
                doc.add_paragraph(text.strip())

    def process_inline_elements(container,paragraph):
        """Process inline elements within a container"""
        for content in container.contents:
            if content.name == 'br':
                paragraph.add_run('\n')
            elif content.name == 'strong':
                run = paragraph.add_run(content.get_text())
                run.bold = True
            elif content.name == 'em':
                run = paragraph.add_run(content.get_text())
                run.italic = True
            elif content.name == 'ul':
                # For nested lists
                for li in content.find_all('li',recursive=False):
                    p = doc.add_paragraph(style='List Bullet')
                    process_inline_elements(li,p)
            elif content.string:
                text = content.string.strip()
                if text:
                    paragraph.add_run(text)

    # Process all top-level elements
    for element in soup.children:
        if element.name:  # Skip NavigableString objects
            process_element(element)

    # Save the document
    doc.save(filename)
    return filename


import random


@app.post("/html_to_word")
async def convert_html_to_word(content: HTMLContent):
    try:
        # Convert HTML to DOCX using htmldocx
        html_content = content.html
        html_content = html_content.replace('\\"','"')  # Replace escaped quotes
        html_content = html_content.replace('\\n','\n')  # Replace escaped newlines
        html_content = html_content.replace('\\r','')  # Remove carriage returns
        print("Contenuto HTML:",html_content)

        parser = HtmlToDocx()
        docx = parser.parse_html_string(html_content)
        filename = "temp" + str(random.choice(range(1000,9999))) + ".docx"
        print("Filename tmp",filename)
        docx.save(filename)
        # Save the DOCX data to a BytesIO object
        with open(filename,mode="rb") as docxfile:
            word_io = docxfile.read()

        os.remove(filename)

        html_markdown_word(html_content,filename)

        with open(filename,mode="rb") as docxfile:
            word_io = docxfile.read()

        os.remove(filename)

        # lines = markdown_text.split('\n')
        return StreamingResponse(io.BytesIO(word_io),
                                 media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                                 headers={"Content-Disposition": "attachment; filename=output.docx"})
    except Exception as e:
        raise HTTPException(status_code=500,detail=str(e))


@app.get("/health")
async def health_check():
    return {"status": "healthy"}